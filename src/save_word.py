# Для работы с docx файлами
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT

# Класс таблиц библиотеки camelot
from camelot.core import Table

# Для проверки существования директорий
import os

# Собственные классы с элементами и шрифтом
from Classes.Font import Font
from Classes.Elements import *

from Tools.config import DEFAULT_TABLE_FONT_SIZE, PAGE_WIDTH, PAGE_HEIGHT


def write_to_word(content_per_page, word_path):
    """
    Записывает содержание каждой страницы в docx файл.

    Args:
        content_per_page: Список словарей, содержащие извлеченное содержимое каждой страницы.
        word_path: Путь к docx файлу.
    """
    
    doc = Document()    # Создаем документ
    first_page = True   # Для проверки на пустую первую секцию
    fonts = []          # Список шрифтов для нахождения прошлых шрифтов
    
    # Цикл по страницам
    for page in content_per_page:   
        # При создании документа создается секция, которую мы получаем
        if first_page:
            section = doc.sections[0]  # Используем первый существующий раздел
            first_page = False
        else:
            section = doc.add_section()  # Создаем новый раздел для следующих страниц
            
        # Изменение ориентации страницы, если альбомная
        set_orientation(section, page['landscape_orientation'])
        
        # Цикл по элементам
        for element in page['page_content']:
            add_text_element(doc, element, fonts)
            add_table_element(doc, element, fonts)

    # Получение директории из пути к файлу
    directory = os.path.dirname(word_path)
                        
    # Проверка на существование директории
    if not os.path.exists(directory):
        os.makedirs(directory)
        
    # Сохранение файла
    doc.save(word_path)
    print(f"The document has been successfully saved: '{word_path}'")
    
# Добавляет Table элемент в документ
def add_table_element(doc, element, fonts):
    """
    Добавляет к документу таблицу из табличного элемента.
    
    Args:
        doc (docx.Document): Объект документа docx, в который будет добавлена таблица.
        element (camelot.core.Table): Объект таблицы, из которой будут извлечены данные.
        fonts (list of Classes.Font): Список, который содержит шрифты других текстовых элементов.
    """
    # Если элемент существует и он является табличным
    if element and isinstance(element, Table):
        # Преобразуем данные Camelot в DataFrame
        table_df = element.df
        
        # Добавляем пустую таблицу в документ
        table = doc.add_table(rows=len(table_df), cols=len(table_df.columns))
        
        # Заполняем таблицу данными
        for i, row in table_df.iterrows():
            # Получаем ячейки строки в таблице
            row_cells = table.rows[i].cells
            for j, item in enumerate(row):
                paragraph = row_cells[j].paragraphs[0]
                run = paragraph.add_run(str(item))
                
                # Меняем стили текста
                font = Font.get_default_font(fonts)
                                        
                run.font.name = font.name
                run.font.size = Pt(font.size) if font.size <= DEFAULT_TABLE_FONT_SIZE else Pt(DEFAULT_TABLE_FONT_SIZE)

# Добавляет TextElement или ImageElement в документ
def add_text_element(doc, element, fonts):
    """
    Добавляет к документу данные из TextElement или ImageElement.
    
    Args:
        doc (docx.Document): Объект документа docx, в который будет добавлена таблица.
        element (Classes.Elements): Объект таблицы, из которой будут извлечены данные.
        fonts (list of Classes.Font): Список, который содержит шрифты других текстовых элементов.
    """
    # Если элемент существует и он является либо текстовым, либо изображением
    if element and (isinstance(element, TextElement) or isinstance(element, ImageElement)):
        # Проверка на пустое содержание текста
        if (element.text.lower().strip() != ""):
            # Создаем элемент параграфа
            paragraph = doc.add_paragraph()
            
            # Создаем элемент Run с текстом
            run = paragraph.add_run(element.text)
            
            # Меняем стили текста
            font = element.font
            if (font):                      
                run.font.name = font.name
                run.font.size = Pt(font.size)
                run.font.bold = font.bold
                run.font.italic = font.italic
                
                fonts.append(font)

def set_orientation(section, is_landscape):
    """
    Меняет ориентацию страницы.

    Args:
    section: Объект секции docx.
    is_landscape: Булево значение, True если нужна альбомная ориентация, иначе False.
    """
    
    if is_landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(PAGE_HEIGHT)
        section.page_height = Inches(PAGE_WIDTH)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(PAGE_WIDTH)
        section.page_height = Inches(PAGE_HEIGHT)
