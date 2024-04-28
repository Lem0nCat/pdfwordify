# To work with docx files
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT

# Camelot library table class
from camelot.core import Table

# To check the existence of directories
import os

# Own classes with elements and font
from pdfwordify.Classes.Font import Font
from pdfwordify.Classes.Elements import *

# Variables from config file
from pdfwordify.Tools.config import DEFAULT_TABLE_FONT_SIZE, PAGE_WIDTH, PAGE_HEIGHT


def write_to_word(content_per_page, word_path):
    """
    Writes the contents of each page to a DOCX file, creating a new section for each page.

    Args:
        content_per_page (list of dicts): A list where each dictionary contains extracted contents of a page.
            Expected keys:
                - 'landscape_orientation' (bool): Indicates if the page should be in landscape orientation.
                - 'page_content' (list): A list of content items where each item can be an item object (pdfwordify.Classes.Elements) or a table (camelot.Table).
        word_path (str): The file path where the DOCX document will be saved.
    """
    
    doc = Document()    # Create a new document
    first_page = True   # Check for an empty first section
    fonts = []          # List of fonts for tracking font styles
    
    # Loop through each page
    for page in content_per_page:
        # The document creation includes one section by default
        if first_page:
            section = doc.sections[0]  # Use the first existing section
            first_page = False
        else:
            section = doc.add_section()  # Create a new section for subsequent pages
            
        # Set page orientation if it's landscape
        set_orientation(section, page['landscape_orientation'])
        
        # Loop through elements on the page
        for element in page['page_content']:
            add_text_element(doc, element, fonts)
            add_table_element(doc, element, fonts)

    # Get the directory from the file path
    directory = os.path.dirname(word_path)
                        
    # Check and create directory if it does not exist
    if not os.path.exists(directory):
        os.makedirs(directory)
        
    # Save the document
    doc.save(word_path)
    
def process_paths(pdf_path, output_path=None):
    """
    Processes input paths for PDF and DOCX files.

    Args:
        pdf_path (str): Path to the PDF file.
        output_path (str, optional): Path where the DOCX file should be saved. Default is None

    Returns:
        str: Corrected output path with .docx extension.
    """
    # Check if output_path is a directory
    if output_path is None:
        # If the path to docx is not specified, create the path by replacing the extension
        docx_path = os.path.splitext(pdf_path)[0] + '.docx'
    elif os.path.isdir(output_path):
        # If the output file is specified as a directory, add the file name and extension
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        docx_path = os.path.join(output_path, f"{base_name}.docx")
    elif not output_path.endswith('.docx'):
        # If a file directory without extension is specified, specify it
        docx_path = f"{output_path}.docx"
    else:
        # If everything is correct
        docx_path = output_path

    return docx_path
    
def add_table_element(doc, element, fonts):
    """
    Adds a table from a table element to the document.
    
    Args:
        doc (docx.Document): The docx document object to which the table will be added.
        element (camelot.core.Table): The object of the table from which the data will be extracted.
        fonts (list of Classes.Font): A list that contains the fonts of other text elements.
    """
    # If the element exists and it is tabular
    if element and isinstance(element, Table):
        # Convert Camelot data into a DataFrame
        table_df = element.df
        
        # Add a blank table to the document
        table = doc.add_table(rows=len(table_df), cols=len(table_df.columns))
        
        # Fill the table with data
        for i, row in table_df.iterrows():
            # Get the row cells in the table
            row_cells = table.rows[i].cells
            for j, item in enumerate(row):
                paragraph = row_cells[j].paragraphs[0]
                run = paragraph.add_run(str(item))
                
                # Change the text styles
                font = Font.get_default_font(fonts)
                                        
                run.font.name = font.name
                run.font.size = Pt(font.size) if font.size <= DEFAULT_TABLE_FONT_SIZE else Pt(DEFAULT_TABLE_FONT_SIZE)

def add_text_element(doc, element, fonts):
    """
    Adds data from a TextElement or ImageElement to the document.
    
    Args:
        doc (docx.Document): The docx document object to which the table will be added.
        element (Classes.Elements): The object of the table from which the data will be extracted.
        fonts (list of Classes.Font): A list that contains the fonts of other text elements.
    """
    # If the element exists and it is either text or image
    if element and (isinstance(element, TextElement) or isinstance(element, ImageElement)):
        # Check for empty text content
        if (element.text.lower().strip() != ""):
            # Create a paragraph element
            paragraph = doc.add_paragraph()
            
            # Create a Run element with text
            run = paragraph.add_run(element.text)
            
            # Change the text styles
            font = element.font
            if (font):                      
                run.font.name = font.name
                run.font.size = Pt(font.size)
                run.font.bold = font.bold
                run.font.italic = font.italic
                
                fonts.append(font)

def set_orientation(section, is_landscape):
    """
    Changes the orientation of the page.

    Args:
    section: docx section object.
    is_landscape: Boolean value, True if landscape orientation is desired, otherwise False.
    """
    
    if is_landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(PAGE_HEIGHT)
        section.page_height = Inches(PAGE_WIDTH)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(PAGE_WIDTH)
        section.page_height = Inches(PAGE_HEIGHT)
