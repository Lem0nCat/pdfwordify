from docx import Document
from docx.shared import Pt

from Font import Font
from Elements import *


# Функция, которая записывает текст в файл Word
def write_to_word_v2(content_per_page, word_path):
    doc = Document()    # Создаем документ

    # Инициализация до цикла, для того, чтобы использовался прошлый шрифт
    font = Font()
    # Цикл по страницам
    for page in content_per_page:
        # Цикл по элементам
        for element in page:
            if element and isinstance(element, TextElement):
                # Проверка на пустое содержание текста
                if (element.text.lower().strip() != ""):
                    # Создаем элемент параграфа
                    paragraph = doc.add_paragraph()
                    
                    # Создаем элемент Run с текстом
                    run = paragraph.add_run(element.text)
                    
                    font = element.font
                    # Меняем стили текста
                    if (font):                      
                        run.font.name = font.name
                        run.font.size = Pt(font.size)
                        run.font.bold = font.bold
                        run.font.italic = font.italic
                    

    doc.save(word_path)
    print(f"Документ успешно сохранен ({word_path})")